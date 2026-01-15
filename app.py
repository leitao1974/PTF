import streamlit as st
from docx import Document
from docx.shared import RGBColor
import google.generativeai as genai
import pandas as pd
import json
import re
from fpdf import FPDF
import pypdf
import io
import time

# ==========================================
# 1. BIBLIOTECA RJAIA (Base de Conhecimento)
# ==========================================
RJAIA_LIBRARY = {
    "Regime Geral": "DL 151-B/2013 alterado pelo DL 11/2023 (Simplex)",
    "Taxas": "Portaria n.º 332-B/2015 (atenção a valores antigos)",
    "Clima": "Lei de Bases do Clima (Lei 98/2021) - obrigatório referir",
    "Prazos": "Verificar deferimentos tácitos do Simplex (DL 11/2023)",
    "Tipologias": "Verificar anexos do DL 151-B/2013 para enquadramento correto"
}

# ==========================================
# 2. FUNÇÕES DE LEITURA E PROCESSAMENTO
# ==========================================

def read_pdf_with_pages(file):
    """Lê PDF e insere marcadores de página."""
    try:
        reader = pypdf.PdfReader(file)
        full_text = ""
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                full_text += f"\n<<<PÁGINA {i+1}>>>\n{text}"
        return full_text
    except Exception as e:
        return f"Erro PDF: {e}"

def read_docx(file):
    """Lê Word e insere marcadores aproximados."""
    doc = Document(file)
    text = ""
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            if i % 20 == 0:
                text += f"\n<<<PARÁGRAFO APROX. {i}>>>\n"
            text += f"{para.text}\n"
    return text

def split_text_into_chunks(text, max_chars=12000):
    """Divide o texto em blocos para não exceder limites da API."""
    chunks = []
    current_chunk = ""
    paragraphs = text.split('\n')
    
    for para in paragraphs:
        if len(current_chunk) + len(para) < max_chars:
            current_chunk += para + "\n"
        else:
            chunks.append(current_chunk)
            current_chunk = para + "\n"
    if current_chunk:
        chunks.append(current_chunk)
    return chunks

def repair_json(json_str):
    """Tenta consertar JSON quebrado (comum em respostas longas de LLMs)."""
    json_str = json_str.strip()
    # Remove formatação markdown se existir
    json_str = re.sub(r"```json\s*|```\s*$", "", json_str).strip()
    
    if not json_str.endswith(']'):
        json_str = json_str.rstrip(',').rstrip() 
        if json_str.count('"') % 2 != 0: json_str += '"'
        if json_str.count('{') > json_str.count('}'): json_str += '}'
        json_str += ']'
    return json_str

# ==========================================
# 3. GERAÇÃO DE RELATÓRIOS (PDF e WORD)
# ==========================================

class PDFReport(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, 'Relatorio Auditoria PTF - RJAIA', 0, 1, 'C')
        self.ln(5)
    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Pag. {self.page_no()}', 0, 0, 'C')

def create_pdf_audit(df):
    """Gera o PDF com a lista de erros."""
    pdf = PDFReport()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    def safe(txt): return str(txt).encode('latin-1', 'replace').decode('latin-1')
    
    for index, row in df.iterrows():
        pdf.set_font('Arial', 'B', 10)
        # Proteção contra Nulos
        loc = safe(row.get('localizacao', '-') or '-')
        cat = safe(row.get('categoria', 'Geral') or 'Geral')
        
        pdf.cell(0, 6, f"Local: {loc} | Tipo: {cat}", 0, 1)
        
        pdf.set_font('Arial', '', 9)
        orig_txt = safe(row.get('texto_detetado', '') or '')
        pdf.multi_cell(0, 5, f"Orig: {orig_txt}")
        
        pdf.set_text_color(200, 0, 0)
        sug_txt = safe(row.get('sugestao', '') or '')
        pdf.multi_cell(0, 5, f"Sug: {sug_txt}")
        pdf.set_text_color(0, 0, 0)
        
        pdf.ln(2)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(2)
    return pdf.output(dest='S').encode('latin-1')

def generate_corrected_docx(original_text, corrections_df):
    """
    Recria o documento Word e aplica correções a vermelho.
    Versão BLINDADA contra erros de NoneType.
    """
    doc = Document()
    doc.add_heading('PTF - Versão Corrigida (IA)', 0)
    
    # Remover marcadores internos (<<<PÁGINA X>>>) para o texto final limpo
    clean_text = re.sub(r'<<<.*?>>>', '', original_text)
    paragraphs = clean_text.split('\n')
    
    # Converter DF para lista de dicionários
    errors = corrections_df.to_dict('records')

    for paragraph in paragraphs:
        if not paragraph.strip(): continue
        p = doc.add_paragraph()
        
        matches = []
        for error in errors:
            # --- CORREÇÃO DO ERRO ---
            # Garante que é string, mesmo que venha None do JSON
            bad = str(error.get('texto_detetado', '') or '').strip()
            good = str(error.get('sugestao', '') or '').strip()
            
            # Só considera se o erro tiver mais de 4 caracteres (evita substituir letras soltas)
            # e se realmente existir no parágrafo
            if len(bad) > 4 and bad in paragraph:
                matches.append((bad, good))
        
        if not matches:
            p.add_run(paragraph)
        else:
            # Se houver erros, aplica a correção no primeiro encontrado (estratégia simples)
            # Ordena por tamanho para garantir que apanha frases completas antes de palavras
            matches.sort(key=lambda x: len(x[0]), reverse=True)
            bad, good = matches[0]
            
            parts = paragraph.split(bad)
            if len(parts) > 1:
                p.add_run(parts[0]) # Texto antes
                
                run_err = p.add_run(good) # Texto corrigido
                run_err.font.color.rgb = RGBColor(255, 0, 0) # Vermelho
                run_err.bold = True
                
                # Texto depois (junta o resto caso o erro apareça mais que uma vez ou o split divida mais)
                p.add_run("".join(parts[1:]))
            else:
                p.add_run(paragraph)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 4. LÓGICA AI (Dinâmica)
# ==========================================

def analyze_chunk(chunk_text, api_key, model_name, library_context):
    """Envia um pedaço do texto para o modelo selecionado."""
    genai.configure(api_key=api_key)
    
    system_prompt = f"""
    És um Especialista em RJAIA (Avaliação de Impacte Ambiental).
    BIBLIOTECA LEGAL: {library_context}
    
    TAREFA: Analisa o texto fornecido e gera um JSON com erros.
    1. Gralhas e Ortografia.
    2. Sintaxe (frases confusas, mal construídas ou pouco técnicas).
    3. Legal (referências erradas ou falta do Simplex DL 11/2023).
    
    OUTPUT JSON (Lista de objetos):
    [
      {{
        "localizacao": "Página X" (ou Parágrafo),
        "categoria": "Sintaxe" (ou Legislação, Gralha),
        "gravidade": "Alta/Baixa",
        "texto_detetado": "trecho exato do texto original com erro",
        "sugestao": "sugestão de correção completa"
      }}
    ]
    """
    
    config = {
        "temperature": 0.1, 
        "response_mime_type": "application/json",
        "max_output_tokens": 8192
    }
    
    try:
        model = genai.GenerativeModel(model_name=model_name, generation_config=config, system_instruction=system_prompt)
        response = model.generate_content(f"Analisa este trecho:\n{chunk_text}")
        return response.text
    except Exception as e:
        return f"ERROR: {str(e)}"

# ==========================================
# 5. INTERFACE (FRONTEND)
# ==========================================

st.set_page_config(page_title="RJAIA Lab", page_icon="🇵🇹", layout="wide")

st.sidebar.header("🔧 Configuração")

# 1. Input da API Key
api_key = st.sidebar.text_input("Google API Key", type="password")

# 2. Deteção Dinâmica de Modelos
available_models = []
if api_key:
    try:
        genai.configure(api_key=api_key)
        # Lista modelos e filtra apenas os que geram texto ('generateContent')
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                available_models.append(m.name)
        
        # Ordena para os mais recentes aparecerem primeiro
        available_models.sort(reverse=True)
        st.sidebar.success(f"Ligação OK! {len(available_models)} modelos detetados.")
        
    except Exception as e:
        st.sidebar.error(f"Erro na chave: {e}")

# Fallback se não detetar
if not available_models:
    available_models = ["models/gemini-1.5-flash", "models/gemini-1.5-pro"]

# 3. Dropdown de Seleção
selected_model = st.sidebar.selectbox("Escolha o Modelo", available_models, index=0)


st.title("🇵🇹 Analisador PTF - RJAIA (Beta)")
st.markdown(f"**Status:** Modelo ativo `{selected_model}`")

uploaded_file = st.file_uploader("Carregue o PTF (Word ou PDF)", type=["docx", "pdf"])

if uploaded_file and api_key:
    if st.button("🚀 Iniciar Análise", type="primary"):
        
        # --- Leitura ---
        fname = uploaded_file.name.lower()
        full_text = ""
        
        if fname.endswith('.pdf'):
            full_text = read_pdf_with_pages(uploaded_file)
        else:
            full_text = read_docx(uploaded_file)
            
        # Validação de Scan
        if len(full_text) < 50:
            st.error("⚠️ Texto demasiado curto ou ilegível.")
            st.warning("Se carregou um PDF digitalizado (imagem), a App não consegue ler. Use um PDF nativo ou Word.")
            st.stop()
            
        # --- Chunking ---
        chunks = split_text_into_chunks(full_text, max_chars=12000)
        st.info(f"Documento processado em {len(chunks)} blocos.")
        
        progress_bar = st.progress(0)
        master_results = []
        library_json = json.dumps(RJAIA_LIBRARY, ensure_ascii=False)
        
        # --- Loop de Processamento ---
        for i, chunk in enumerate(chunks):
            # Mensagem de estado
            with st.spinner(f"A analisar parte {i+1} de {len(chunks)}..."):
                raw_resp = analyze_chunk(chunk, api_key, selected_model, library_json)
                
                if not raw_resp.startswith("ERROR"):
                    try:
                        # Limpeza
                        cleaned = repair_json(raw_resp)
                        data = json.loads(cleaned)
                        
                        # Normalização (garantir lista)
                        if isinstance(data, dict): data = [data]
                        if isinstance(data, list): master_results.extend(data)
                    except Exception as e:
                        # Log discreto se falhar um chunk, não para o processo
                        print(f"Erro parse chunk {i}: {e}")
                else:
                    st.warning(f"Erro na API (bloco {i+1}): {raw_resp}")
            
            # Atualiza progresso
            progress_bar.progress((i + 1) / len(chunks))
            # Pausa curta para evitar rate-limit se usar chave gratuita
            time.sleep(1) 
            
        st.success("Análise completa!")
        
        # Guardar em Sessão (para não perder ao clicar nos downloads)
        st.session_state['results'] = pd.DataFrame(master_results)
        st.session_state['text_ref'] = full_text

# --- Exibição de Resultados ---
if 'results' in st.session_state:
    df = st.session_state['results']
    
    if not df.empty:
        col1, col2 = st.columns([1, 2])
        
        with col1:
            st.markdown("### Resumo")
            st.metric("Total de Observações", len(df))
            st.metric("Erros Legais", len(df[df['categoria'].str.contains('Legisl', case=False, na=False)]))
            
            st.divider()
            st.markdown("### 📥 Downloads")
            
            # Download PDF
            pdf_out = create_pdf_audit(df)
            st.download_button(
                "📄 Relatório de Auditoria (PDF)", 
                pdf_out, 
                "relatorio_auditoria.pdf", 
                "application/pdf"
            )
            
            # Download Word
            doc_out = generate_corrected_docx(st.session_state['text_ref'], df)
            st.download_button(
                "📝 PTF com Correções (Word)", 
                doc_out, 
                "ptf_corrigido.docx", 
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        with col2:
            st.markdown("### Detalhes")
            st.dataframe(
                df, 
                column_config={
                    "localizacao": "Local",
                    "categoria": "Tipo",
                    "texto_detetado": "Original",
                    "sugestao": "Sugestão"
                },
                use_container_width=True,
                hide_index=True
            )
    else:
        st.warning("O modelo analisou o documento mas não reportou erros. (Ou houve falha na interpretação da resposta).")

elif not api_key:
    st.info("👈 Insira a Google API Key na barra lateral para começar.")
